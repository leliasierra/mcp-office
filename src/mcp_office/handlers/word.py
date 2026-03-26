import json
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from mcp_office.handlers.base import DocumentHandler
from mcp.types import Tool, TextContent


class WordHandler(DocumentHandler):
    """Handler for Word document operations - Single Responsibility"""

    def get_tools(self) -> list[Tool]:
        return [
            Tool(
                name="create_document",
                description="Create a new Word document",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {
                            "type": "string",
                            "description": "Path to save document",
                        },
                        "title": {"type": "string", "description": "Document title"},
                        "author": {"type": "string", "description": "Document author"},
                    },
                    "required": ["path"],
                },
            ),
            Tool(
                name="open_document",
                description="Open and analyze a Word document",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Path to document"},
                        "extract_text": {
                            "type": "boolean",
                            "description": "Extract text",
                        },
                    },
                    "required": ["path"],
                },
            ),
            Tool(
                name="get_document_properties",
                description="Get document properties",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Path to document"},
                    },
                    "required": ["path"],
                },
            ),
            Tool(
                name="list_documents",
                description="List documents in directory",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "directory": {
                            "type": "string",
                            "description": "Directory path",
                        },
                        "pattern": {"type": "string", "description": "File pattern"},
                    },
                    "required": ["directory"],
                },
            ),
            Tool(
                name="copy_document",
                description="Copy a document",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "source": {"type": "string", "description": "Source path"},
                        "destination": {
                            "type": "string",
                            "description": "Destination path",
                        },
                    },
                    "required": ["source", "destination"],
                },
            ),
            Tool(
                name="merge_documents",
                description="Merge multiple documents",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "sources": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "Source document paths",
                        },
                        "output": {"type": "string", "description": "Output path"},
                    },
                    "required": ["sources", "output"],
                },
            ),
            Tool(
                name="add_heading",
                description="Add a heading",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Document path"},
                        "text": {"type": "string", "description": "Heading text"},
                        "level": {"type": "integer", "description": "Heading level"},
                        "bold": {"type": "boolean", "description": "Bold text"},
                    },
                    "required": ["path", "text", "level"],
                },
            ),
            Tool(
                name="add_paragraph",
                description="Add a paragraph",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Document path"},
                        "text": {"type": "string", "description": "Paragraph text"},
                        "bold": {"type": "boolean", "description": "Bold text"},
                        "alignment": {"type": "string", "description": "Alignment"},
                    },
                    "required": ["path", "text"],
                },
            ),
            Tool(
                name="add_table",
                description="Add a table",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Document path"},
                        "data": {
                            "type": "array",
                            "items": {"type": "array", "items": {"type": "string"}},
                            "description": "Table data",
                        },
                        "header_row": {"type": "boolean", "description": "Header row"},
                    },
                    "required": ["path", "data"],
                },
            ),
            Tool(
                name="add_image",
                description="Add an image",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Document path"},
                        "image_path": {"type": "string", "description": "Image path"},
                        "width_cm": {"type": "number", "description": "Width in cm"},
                    },
                    "required": ["path", "image_path"],
                },
            ),
            Tool(
                name="add_page_break",
                description="Add a page break",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Document path"},
                    },
                    "required": ["path"],
                },
            ),
            Tool(
                name="add_bullet_list",
                description="Add a bullet/numbered list",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Document path"},
                        "items": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "List items",
                        },
                        "numbered": {"type": "boolean", "description": "Numbered list"},
                    },
                    "required": ["path", "items"],
                },
            ),
            Tool(
                name="find_replace",
                description="Find and replace text",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Document path"},
                        "find": {"type": "string", "description": "Text to find"},
                        "replace": {
                            "type": "string",
                            "description": "Replacement text",
                        },
                    },
                    "required": ["path", "find", "replace"],
                },
            ),
            Tool(
                name="create_from_spec",
                description="Create document from JSON specification",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "spec_path": {
                            "type": "string",
                            "description": "JSON spec path",
                        },
                        "output_path": {
                            "type": "string",
                            "description": "Output document path",
                        },
                    },
                    "required": ["spec_path", "output_path"],
                },
            ),
            Tool(
                name="save_document_spec",
                description="Save document specification to JSON for reuse",
                inputSchema={
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Document path"},
                        "output_path": {
                            "type": "string",
                            "description": "JSON output path",
                        },
                    },
                    "required": ["path", "output_path"],
                },
            ),
        ]

    async def execute(self, tool_name: str, arguments: dict) -> list[TextContent]:
        try:
            if tool_name == "create_document":
                return self._create_document(arguments)
            elif tool_name == "open_document":
                return self._open_document(arguments)
            elif tool_name == "get_document_properties":
                return self._get_properties(arguments)
            elif tool_name == "list_documents":
                return self._list_documents(arguments)
            elif tool_name == "copy_document":
                return self._copy_document(arguments)
            elif tool_name == "merge_documents":
                return self._merge_documents(arguments)
            elif tool_name == "add_heading":
                return self._add_heading(arguments)
            elif tool_name == "add_paragraph":
                return self._add_paragraph(arguments)
            elif tool_name == "add_table":
                return self._add_table(arguments)
            elif tool_name == "add_image":
                return self._add_image(arguments)
            elif tool_name == "add_page_break":
                return self._add_page_break(arguments)
            elif tool_name == "add_bullet_list":
                return self._add_bullet_list(arguments)
            elif tool_name == "find_replace":
                return self._find_replace(arguments)
            elif tool_name == "format_text_range":
                return self._format_text_range(arguments)
            elif tool_name == "protect_document":
                return self._protect_document(arguments)
            elif tool_name == "extract_comments":
                return self._extract_comments(arguments)
            elif tool_name == "create_from_spec":
                return self._create_from_spec(arguments)
            elif tool_name == "save_document_spec":
                return self._save_document_spec(arguments)
            return self.error_result(f"Unknown tool: {tool_name}")
        except Exception as e:
            return self.error_result(str(e))

    def _create_document(self, args: dict) -> list[TextContent]:
        doc = Document()
        if title := args.get("title"):
            doc.core_properties.title = title
        if author := args.get("author"):
            doc.core_properties.author = author
        doc.save(args["path"])
        return self.success_result(f"Document created: {args['path']}")

    def _open_document(self, args: dict) -> list[TextContent]:
        doc = Document(args["path"])
        result = {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}
        if args.get("extract_text"):
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            result["text"] = text[:5000]
        return [TextContent(type="text", text=json.dumps(result, indent=2))]

    def _get_properties(self, args: dict) -> list[TextContent]:
        doc = Document(args["path"])
        props = doc.core_properties
        return [
            TextContent(
                type="text",
                text=json.dumps(
                    {
                        "title": props.title or "",
                        "author": props.author or "",
                        "paragraphs": len(doc.paragraphs),
                        "tables": len(doc.tables),
                    },
                    indent=2,
                ),
            )
        ]

    def _list_documents(self, args: dict) -> list[TextContent]:
        directory = args.get("directory", ".")
        pattern = args.get("pattern", "*.docx")
        docs = list(Path(directory).glob(pattern))
        return [
            TextContent(type="text", text=json.dumps([str(d) for d in docs], indent=2))
        ]

    def _copy_document(self, args: dict) -> list[TextContent]:
        shutil.copy2(args["source"], args["destination"])
        return self.success_result(f"Document copied to: {args['destination']}")

    def _merge_documents(self, args: dict) -> list[TextContent]:
        merged = Document()
        for src_path in args["sources"]:
            src_doc = Document(src_path)
            for elem in src_doc.element.body:
                merged.element.body.append(elem)
        merged.save(args["output"])
        return self.success_result(f"Documents merged: {args['output']}")

    def _add_heading(self, args: dict) -> list[TextContent]:
        doc = Document(args["path"])
        heading = doc.add_heading(args["text"], args.get("level", 1))
        if args.get("bold"):
            heading.runs[0].bold = True
        doc.save(args["path"])
        return self.success_result(f"Heading added: {args['text']}")

    def _add_paragraph(self, args: dict) -> list[TextContent]:
        doc = Document(args["path"])
        paragraph = doc.add_paragraph(args["text"])
        if args.get("bold"):
            for run in paragraph.runs:
                run.bold = True
        if alignment := args.get("alignment"):
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }
            paragraph.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        doc.save(args["path"])
        return self.success_result(f"Paragraph added: {args['text']}")

    def _add_table(self, args: dict) -> list[TextContent]:
        doc = Document(args["path"])
        data = args["data"]
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        for i, row_data in enumerate(data):
            for j, cell_text in enumerate(row_data):
                table.rows[i].cells[j].text = cell_text
        if args.get("header_row"):
            for cell in table.rows[0].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
        doc.save(args["path"])
        return self.success_result(f"Table added with {len(data)} rows")

    def _add_image(self, args: dict) -> list[TextContent]:
        doc = Document(args["path"])
        width = Inches(args.get("width_cm", 5)) if args.get("width_cm") else None
        doc.add_picture(args["image_path"], width=width)
        doc.save(args["path"])
        return self.success_result(f"Image added: {args['image_path']}")

    def _add_page_break(self, args: dict) -> list[TextContent]:
        doc = Document(args["path"])
        doc.add_page_break()
        doc.save(args["path"])
        return self.success_result("Page break added")

    def _add_bullet_list(self, args: dict) -> list[TextContent]:
        doc = Document(args["path"])
        if args.get("numbered"):
            for i, item in enumerate(args["items"], start=1):
                doc.add_paragraph(f"{i}. {item}", style="List Number")
        else:
            for item in args["items"]:
                doc.add_paragraph(item, style="List Bullet")
        doc.save(args["path"])
        return self.success_result(f"List added with {len(args['items'])} items")

    def _find_replace(self, args: dict) -> list[TextContent]:
        doc = Document(args["path"])
        count = 0
        for para in doc.paragraphs:
            if args["find"] in para.text:
                para.text = para.text.replace(args["find"], args["replace"])
                count += 1
        doc.save(args["path"])
        return self.success_result(f"Replaced {count} occurrences")

    def _format_text_range(self, args: dict) -> list[TextContent]:
        doc = Document(args["path"])
        for para in doc.paragraphs:
            if args["start_text"] in para.text:
                for run in para.runs:
                    if args["start_text"] in run.text:
                        if args.get("bold"):
                            run.bold = True
        doc.save(args["path"])
        return self.success_result("Text formatted")

    def _protect_document(self, args: dict) -> list[TextContent]:
        doc = Document(args["path"])
        doc.settings.password = args["password"]
        doc.save(args["path"])
        return self.success_result("Document protected with password")

    def _extract_comments(self, args: dict) -> list[TextContent]:
        doc = Document(args["path"])
        comments = []
        try:
            for comment in doc.part.comments.comments:
                comments.append(
                    {"author": comment.author or "", "text": comment.text or ""}
                )
        except AttributeError:
            pass
        return [TextContent(type="text", text=json.dumps(comments, indent=2))]

    def _create_from_spec(self, args: dict) -> list[TextContent]:
        with open(args["spec_path"]) as f:
            spec = json.load(f)

        doc = Document()

        for item in spec.get("content", []):
            if item["type"] == "heading":
                doc.add_heading(item["text"], item.get("level", 1))
            elif item["type"] == "paragraph":
                p = doc.add_paragraph(item.get("text", ""))
                if item.get("bold"):
                    for run in p.runs:
                        run.bold = True
            elif item["type"] == "table":
                table = doc.add_table(rows=len(item["data"]), cols=len(item["data"][0]))
                for i, row in enumerate(item["data"]):
                    for j, cell in enumerate(row):
                        table.rows[i].cells[j].text = cell
            elif item["type"] == "image":
                width = Inches(item.get("width", 5))
                doc.add_picture(item["path"], width=width)
            elif item["type"] == "page_break":
                doc.add_page_break()
            elif item["type"] == "list":
                style = "List Number" if item.get("numbered") else "List Bullet"
                for list_item in item.get("items", []):
                    doc.add_paragraph(list_item, style=style)

        output_path = args.get(
            "output_path", args["spec_path"].replace(".json", ".docx")
        )
        doc.save(output_path)

        return self.success_result(f"Document created from spec: {output_path}")

    def _save_document_spec(self, args: dict) -> list[TextContent]:
        doc = Document(args["path"])
        spec = {"content": []}

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                spec["content"].append(
                    {"type": "heading", "text": para.text, "level": level}
                )
            else:
                is_bold = any(run.bold for run in para.runs if run.text)
                spec["content"].append(
                    {"type": "paragraph", "text": para.text, "bold": is_bold}
                )

        for table in doc.tables:
            data = [[cell.text for cell in row.cells] for row in table.rows]
            spec["content"].append({"type": "table", "data": data})

        output_path = args.get(
            "output_path", args["path"].replace(".docx", "_spec.json")
        )
        with open(output_path, "w") as f:
            json.dump(spec, f, indent=2)

        return self.success_result(f"Document spec saved: {output_path}")
