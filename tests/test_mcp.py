import pytest
import json
from pathlib import Path
from docx import Document
from openpyxl import Workbook
from pptx import Presentation
from reportlab.pdfgen import canvas
from mcp_office import list_tools, call_tool


@pytest.fixture
def temp_doc(tmp_path):
    doc_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Test paragraph")
    doc.save(str(doc_path))
    return str(doc_path)


@pytest.fixture
def temp_xlsx(tmp_path):
    xlsx_path = tmp_path / "test.xlsx"
    wb = Workbook()
    wb.active.append(["Header1", "Header2"])
    wb.save(str(xlsx_path))
    return str(xlsx_path)


@pytest.fixture
def temp_pdf(tmp_path):
    pdf_path = tmp_path / "test.pdf"
    return str(pdf_path)


@pytest.mark.asyncio
async def test_list_tools():
    tools = await list_tools()
    assert len(tools) > 0
    tool_names = [t.name for t in tools]
    assert "create_document" in tool_names
    assert "open_document" in tool_names
    assert "pdf_read" in tool_names
    assert "excel_create" in tool_names
    assert "ppt_generate" in tool_names


@pytest.mark.asyncio
async def test_create_document(tmp_path):
    path = str(tmp_path / "new_doc.docx")
    result = await call_tool("create_document", {"path": path})
    assert "created" in result[0].text.lower()
    assert Path(path).exists()


@pytest.mark.asyncio
async def test_open_document(temp_doc):
    result = await call_tool("open_document", {"path": temp_doc, "extract_text": True})
    data = json.loads(result[0].text)
    assert "paragraphs" in data
    assert data["paragraphs"] >= 1


@pytest.mark.asyncio
async def test_list_documents(tmp_path):
    doc_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(str(doc_path))

    result = await call_tool("list_documents", {"directory": str(tmp_path)})
    files = json.loads(result[0].text)
    assert len(files) >= 1
    assert any("test.docx" in f for f in files)


@pytest.mark.asyncio
async def test_get_document_properties(temp_doc):
    result = await call_tool("get_document_properties", {"path": temp_doc})
    data = json.loads(result[0].text)
    assert "paragraphs" in data
    assert "tables" in data


@pytest.mark.asyncio
async def test_copy_document(temp_doc, tmp_path):
    dest = str(tmp_path / "copy.docx")
    result = await call_tool("copy_document", {"source": temp_doc, "destination": dest})
    assert "copied" in result[0].text.lower()
    assert Path(dest).exists()


@pytest.mark.asyncio
async def test_add_paragraph(temp_doc):
    result = await call_tool(
        "add_paragraph", {"path": temp_doc, "text": "New paragraph", "bold": True}
    )
    assert "paragraph" in result[0].text.lower()


@pytest.mark.asyncio
async def test_add_heading(temp_doc):
    result = await call_tool(
        "add_heading",
        {"path": temp_doc, "text": "Test Heading", "level": 1, "bold": True},
    )
    assert "heading" in result[0].text.lower()


@pytest.mark.asyncio
async def test_find_replace(temp_doc):
    await call_tool("add_paragraph", {"path": temp_doc, "text": "Hello World"})
    result = await call_tool(
        "find_replace", {"path": temp_doc, "find": "World", "replace": "Python"}
    )
    assert "replaced" in result[0].text.lower()


@pytest.mark.asyncio
async def test_add_table(temp_doc):
    data = [["Header1", "Header2"], ["Row1Col1", "Row1Col2"]]
    result = await call_tool(
        "add_table", {"path": temp_doc, "data": data, "header_row": True}
    )
    assert "table" in result[0].text.lower()


@pytest.mark.asyncio
async def test_add_page_break(temp_doc):
    result = await call_tool("add_page_break", {"path": temp_doc})
    assert "page" in result[0].text.lower()


@pytest.mark.asyncio
async def test_bullet_list(temp_doc):
    result = await call_tool(
        "add_bullet_list", {"path": temp_doc, "items": ["Item 1", "Item 2", "Item 3"]}
    )
    assert "list" in result[0].text.lower()


@pytest.mark.asyncio
async def test_invalid_tool():
    result = await call_tool("nonexistent_tool", {})
    assert "unknown tool" in result[0].text.lower()


@pytest.mark.asyncio
async def test_excel_create(tmp_path):
    path = str(tmp_path / "new_book.xlsx")
    result = await call_tool("excel_create", {"path": path})
    assert "created" in result[0].text.lower()
    assert Path(path).exists()


@pytest.mark.asyncio
async def test_excel_read(temp_xlsx):
    result = await call_tool("excel_read", {"path": temp_xlsx})
    data = json.loads(result[0].text)
    assert len(data) > 0
    assert "Header1" in data[0]


@pytest.mark.asyncio
async def test_excel_write_cell(temp_xlsx):
    result = await call_tool(
        "excel_write_cell",
        {"path": temp_xlsx, "sheet": "Sheet", "cell": "A2", "value": "TestValue"},
    )
    assert "wrote" in result[0].text.lower()


@pytest.mark.asyncio
async def test_excel_add_row(temp_xlsx):
    result = await call_tool(
        "excel_add_row",
        {"path": temp_xlsx, "sheet": "Sheet", "data": ["Data1", "Data2"]},
    )
    assert "added" in result[0].text.lower()


@pytest.mark.asyncio
async def test_excel_add_formula(temp_xlsx):
    result = await call_tool(
        "excel_add_formula",
        {"path": temp_xlsx, "sheet": "Sheet", "cell": "C1", "formula": "=SUM(A1:B1)"},
    )
    assert "formula" in result[0].text.lower()


@pytest.mark.asyncio
async def test_excel_format_cell(temp_xlsx):
    result = await call_tool(
        "excel_format_cell",
        {
            "path": temp_xlsx,
            "sheet": "Sheet",
            "cell": "A1",
            "bold": True,
            "font_size": 14,
        },
    )
    assert "formatted" in result[0].text.lower()


@pytest.mark.asyncio
async def test_excel_add_table(temp_xlsx):
    result = await call_tool(
        "excel_add_table",
        {
            "path": temp_xlsx,
            "sheet": "Sheet",
            "data": [["A", "B"], ["1", "2"]],
            "start_cell": "E1",
        },
    )
    assert "table" in result[0].text.lower()


@pytest.mark.asyncio
async def test_excel_set_column_width(temp_xlsx):
    result = await call_tool(
        "excel_set_column_width",
        {"path": temp_xlsx, "sheet": "Sheet", "column": "A", "width": 20},
    )
    assert "width" in result[0].text.lower()


@pytest.mark.asyncio
async def test_excel_merge_cells(temp_xlsx):
    result = await call_tool(
        "excel_merge_cells",
        {"path": temp_xlsx, "sheet": "Sheet", "start_cell": "A1", "end_cell": "B1"},
    )
    assert "merged" in result[0].text.lower()


@pytest.mark.asyncio
async def test_excel_list_sheets(temp_xlsx):
    result = await call_tool("excel_list_sheets", {"path": temp_xlsx})
    sheets = json.loads(result[0].text)
    assert len(sheets) >= 1


@pytest.fixture
def temp_svg_dir(tmp_path):
    svg_dir = tmp_path / "slides"
    svg_dir.mkdir()
    svg_1 = svg_dir / "01_cover.svg"
    svg_1.write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1280 720">'
        '<rect width="1280" height="720" fill="#1a1a2e"/>'
        '<text x="640" y="360" text-anchor="middle" font-size="48" fill="white">'
        "Test Title</text></svg>"
    )
    svg_2 = svg_dir / "02_content.svg"
    svg_2.write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1280 720">'
        '<rect width="1280" height="720" fill="#ffffff"/>'
        '<text x="640" y="360" text-anchor="middle" font-size="36" fill="#333">'
        "Test Content</text></svg>"
    )
    return str(svg_dir)


@pytest.mark.asyncio
async def test_ppt_generate(temp_svg_dir, tmp_path):
    output = str(tmp_path / "presentation.pptx")
    result = await call_tool(
        "ppt_generate", {"svg_source": temp_svg_dir, "output": output}
    )
    assert "generated" in result[0].text.lower()
    assert Path(output).exists()


@pytest.mark.asyncio
async def test_ppt_generate_invalid_source(tmp_path):
    result = await call_tool(
        "ppt_generate",
        {"svg_source": str(tmp_path / "nonexistent"), "output": str(tmp_path / "out.pptx")},
    )
    assert "error" in result[0].text.lower()


@pytest.mark.asyncio
async def test_excel_create_xlsxwriter(tmp_path):
    path = str(tmp_path / "xlsxw.xlsx")
    result = await call_tool("excel_create_xlsxwriter", {"path": path})
    assert "created" in result[0].text.lower()
    assert Path(path).exists()


@pytest.mark.asyncio
async def test_excel_add_chart(tmp_path):
    path = str(tmp_path / "chart.xlsx")
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Category"
    ws["B1"] = "Value"
    ws["A2"] = "A"
    ws["B2"] = 10
    ws["A3"] = "B"
    ws["B3"] = 20
    wb.save(path)

    result = await call_tool(
        "excel_add_chart",
        {
            "path": path,
            "chart_type": "column",
            "data_range": "Sheet!$A$1:$B$3",
            "title": "Test Chart",
        },
    )
    assert "chart" in result[0].text.lower()


@pytest.mark.asyncio
async def test_excel_pandas_to_excel(tmp_path):
    path = str(tmp_path / "pandas.xlsx")
    data = json.dumps([{"Name": "Alice", "Age": 30}, {"Name": "Bob", "Age": 25}])
    result = await call_tool(
        "excel_pandas_to_excel", {"path": path, "data": data, "sheet_name": "Data"}
    )
    assert "created" in result[0].text.lower()
    assert Path(path).exists()


@pytest.mark.asyncio
async def test_pdf_create(tmp_path):
    path = str(tmp_path / "new.pdf")
    result = await call_tool("pdf_create", {"path": path, "title": "Test PDF"})
    assert "created" in result[0].text.lower()
    assert Path(path).exists()


@pytest.mark.asyncio
async def test_pdf_add_text(tmp_path):
    pdf_path = tmp_path / "test.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.save()

    result = await call_tool(
        "pdf_add_text",
        {
            "path": str(pdf_path),
            "text": "Hello World",
            "x": 100,
            "y": 700,
            "font_size": 12,
        },
    )
    assert "added" in result[0].text.lower()


@pytest.mark.asyncio
async def test_pdf_protect(tmp_path):
    pdf_path = tmp_path / "protected.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.save()

    result = await call_tool(
        "pdf_protect",
        {"path": str(pdf_path), "password": "test123"},
    )
    assert "protected" in result[0].text.lower()


@pytest.mark.asyncio
async def test_pdf_unprotect(tmp_path):
    pdf_path = tmp_path / "locked.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.save()

    await call_tool("pdf_protect", {"path": str(pdf_path), "password": "test123"})

    unlocked_path = tmp_path / "unlocked.pdf"
    result = await call_tool(
        "pdf_unprotect",
        {"path": str(pdf_path), "password": "test123", "output": str(unlocked_path)},
    )
    assert "unlocked" in result[0].text.lower()
    assert unlocked_path.exists()


@pytest.mark.asyncio
async def test_create_from_spec(tmp_path):
    spec_path = tmp_path / "spec.json"
    spec = {
        "content": [
            {"type": "heading", "text": "Title", "level": 1},
            {"type": "paragraph", "text": "Hello world", "bold": True},
            {"type": "table", "data": [["A", "B"], ["1", "2"]]},
        ]
    }
    with open(spec_path, "w") as f:
        json.dump(spec, f)

    output_path = tmp_path / "output.docx"
    result = await call_tool(
        "create_from_spec",
        {"spec_path": str(spec_path), "output_path": str(output_path)},
    )
    assert "created" in result[0].text.lower()
    assert output_path.exists()


@pytest.mark.asyncio
async def test_save_document_spec(temp_doc, tmp_path):
    output_path = tmp_path / "spec.json"
    result = await call_tool(
        "save_document_spec",
        {"path": temp_doc, "output_path": str(output_path)},
    )
    assert "saved" in result[0].text.lower()
    assert output_path.exists()


@pytest.mark.asyncio
async def test_auto_spec_word_full_flow(tmp_path):
    doc_path = str(tmp_path / "full_test.docx")
    spec_path = doc_path + ".spec.json"

    # --- Create document with properties ---
    result = await call_tool("create_document", {"path": doc_path, "title": "Full Test", "author": "Tester"})
    assert "created" in result[0].text.lower()
    assert Path(doc_path).exists()
    spec = json.loads(Path(spec_path).read_text())
    assert spec["type"] == "word"
    assert spec["path"] == doc_path
    assert spec["properties"]["title"] == "Full Test"
    assert spec["properties"]["author"] == "Tester"
    assert spec["content"][0]["type"] == "create"
    assert len(spec["content"]) == 1

    # --- Add heading ---
    await call_tool("add_heading", {"path": doc_path, "text": "Introduction", "level": 1, "bold": True})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["content"][-1] == {"type": "heading", "text": "Introduction", "level": 1, "bold": True}
    assert len(spec["content"]) == 2

    # --- Add paragraph with alignment and bold ---
    await call_tool("add_paragraph", {"path": doc_path, "text": "This is a bold center paragraph.", "bold": True, "alignment": "center"})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["content"][-1] == {"type": "paragraph", "text": "This is a bold center paragraph.", "bold": True, "alignment": "center"}
    assert len(spec["content"]) == 3

    # --- Add another paragraph (plain) ---
    await call_tool("add_paragraph", {"path": doc_path, "text": "Plain left text."})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["content"][-1] == {"type": "paragraph", "text": "Plain left text.", "bold": False, "alignment": None}
    assert len(spec["content"]) == 4

    # --- Add table ---
    data = [["Name", "Age"], ["Alice", "30"], ["Bob", "25"]]
    await call_tool("add_table", {"path": doc_path, "data": data, "header_row": True})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["content"][-1]["type"] == "table"
    assert spec["content"][-1]["data"] == data
    assert spec["content"][-1]["header_row"] is True
    assert len(spec["content"]) == 5

    # --- Add bullet list ---
    await call_tool("add_bullet_list", {"path": doc_path, "items": ["First", "Second", "Third"], "numbered": False})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["content"][-1]["type"] == "list"
    assert spec["content"][-1]["items"] == ["First", "Second", "Third"]
    assert spec["content"][-1]["numbered"] is False
    assert len(spec["content"]) == 6

    # --- Add numbered list ---
    await call_tool("add_bullet_list", {"path": doc_path, "items": ["Step 1", "Step 2"], "numbered": True})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["content"][-1]["numbered"] is True
    assert len(spec["content"]) == 7

    # --- Add page break ---
    await call_tool("add_page_break", {"path": doc_path})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["content"][-1] == {"type": "page_break"}
    assert len(spec["content"]) == 8

    # --- Find and replace ---
    await call_tool("find_replace", {"path": doc_path, "find": "Alice", "replace": "Charlie"})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["content"][-1]["type"] == "find_replace"
    assert spec["content"][-1]["find"] == "Alice"
    assert spec["content"][-1]["replace"] == "Charlie"
    assert len(spec["content"]) == 9

    # --- Copy document and verify spec is copied too ---
    copy_path = str(tmp_path / "copy_test.docx")
    await call_tool("copy_document", {"source": doc_path, "destination": copy_path})
    assert Path(copy_path).exists()
    assert Path(copy_path + ".spec.json").exists()
    copy_spec = json.loads(Path(copy_path + ".spec.json").read_text())
    assert copy_spec["content"] == spec["content"]

    # --- Verify content ordering ---
    assert [c["type"] for c in spec["content"]] == [
        "create", "heading", "paragraph", "paragraph",
        "table", "list", "list", "page_break", "find_replace"
    ]


@pytest.mark.asyncio
async def test_auto_spec_excel_full_flow(tmp_path):
    xlsx_path = str(tmp_path / "full_test.xlsx")
    spec_path = xlsx_path + ".spec.json"

    # --- Create workbook ---
    await call_tool("excel_create", {"path": xlsx_path, "sheet_name": "Data"})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["type"] == "excel"
    assert spec["properties"]["sheet_name"] == "Data"
    assert spec["content"][0]["type"] == "create"

    # --- Write cell ---
    await call_tool("excel_write_cell", {"path": xlsx_path, "sheet": "Data", "cell": "A1", "value": "Hello"})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["content"][-1] == {"type": "write_cell", "sheet": "Data", "cell": "A1", "value": "Hello"}

    # --- Add row ---
    await call_tool("excel_add_row", {"path": xlsx_path, "sheet": "Data", "data": ["Alice", "30"]})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["content"][-1]["type"] == "add_row"
    assert spec["content"][-1]["data"] == ["Alice", "30"]

    # --- Add formula ---
    await call_tool("excel_add_formula", {"path": xlsx_path, "sheet": "Data", "cell": "C1", "formula": "=SUM(A1:B1)"})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["content"][-1]["type"] == "add_formula"
    assert spec["content"][-1]["formula"] == "=SUM(A1:B1)"

    # --- Format cell ---
    await call_tool("excel_format_cell", {"path": xlsx_path, "sheet": "Data", "cell": "A1", "bold": True, "font_size": 14, "font_color": "#FF0000"})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["content"][-1]["type"] == "format_cell"
    assert spec["content"][-1]["bold"] is True

    # --- Add table ---
    await call_tool("excel_add_table", {"path": xlsx_path, "sheet": "Data", "data": [["X", "Y"], ["1", "2"]], "start_cell": "E1"})
    spec = json.loads(Path(spec_path).read_text())
    assert spec["content"][-1]["type"] == "add_table"

    # --- Record conversion intent ---
    await call_tool("excel_to_pdf", {"input_path": xlsx_path, "output_path": str(tmp_path / "test.pdf")})
    spec = json.loads(Path(spec_path).read_text())
    assert len(spec["post_processing"]) == 1
    assert spec["post_processing"][0]["action"] == "convert_to_pdf"

    # --- Verify all content types ---
    types = [c["type"] for c in spec["content"]]
    assert types == ["create", "write_cell", "add_row", "add_formula", "format_cell", "add_table"]


@pytest.mark.asyncio
async def test_ppt_generate_auto_spec(temp_svg_dir, tmp_path):
    pptx_path = str(tmp_path / "full_test.pptx")
    spec_path = pptx_path + ".spec.json"

    result = await call_tool(
        "ppt_generate", {"svg_source": temp_svg_dir, "output": pptx_path}
    )
    assert "generated" in result[0].text.lower()
    assert Path(pptx_path).exists()

    spec = json.loads(Path(spec_path).read_text())
    assert spec["type"] == "powerpoint"
    assert spec["content"][0]["type"] == "generate"
    assert spec["content"][0]["slides"] == 2


@pytest.mark.asyncio
async def test_auto_spec_create_from_spec(tmp_path):
    spec_path = str(tmp_path / "my_spec.json")
    spec_data = {
        "type": "word",
        "path": "",
        "properties": {"title": "Generated"},
        "content": [
            {"type": "heading", "text": "Title", "level": 1, "bold": True},
            {"type": "paragraph", "text": "Content", "bold": False, "alignment": None},
        ],
        "post_processing": [],
    }
    with open(spec_path, "w") as f:
        json.dump(spec_data, f)

    output_path = str(tmp_path / "from_spec.docx")
    await call_tool("create_from_spec", {"spec_path": spec_path, "output_path": output_path})
    assert Path(output_path).exists()
    auto_spec = json.loads(Path(output_path + ".spec.json").read_text())
    assert auto_spec["type"] == "word"
    assert auto_spec["path"] == output_path
    assert len(auto_spec["content"]) == 2
    assert auto_spec["content"][0]["type"] == "heading"
